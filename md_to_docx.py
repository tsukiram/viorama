"""Convert markdown to docx with academic formatting.

Usage:
    python3 md_to_docx.py SRC.md DST.docx
"""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

if len(sys.argv) >= 3:
    SRC, DST = sys.argv[1], sys.argv[2]
else:
    SRC = "/Users/macbookpro/Desktop/viorama-website-new/bab2_outline.md"
    DST = "/Users/macbookpro/Desktop/viorama-website-new/bab2.docx"

doc = Document()

# Page margins (Indonesian thesis convention: 4-3-3-3)
section = doc.sections[0]
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin = Cm(4)
section.right_margin = Cm(3)

# Default font: Times New Roman 12 pt, line spacing 1.5
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)


def add_runs(paragraph, text):
    """Parse inline **bold** and *italic* and add runs to paragraph."""
    pos = 0
    pattern = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*([^*]+?)\*")
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        if m.group(1) is not None:
            run = paragraph.add_run(m.group(1))
            run.bold = True
            run.italic = True
        elif m.group(2) is not None:
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif m.group(3) is not None:
            run = paragraph.add_run(m.group(3))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_heading(text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(14)
    elif level == 2:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
    else:
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)


def add_paragraph(text, justify=True, indent_first=True):
    p = doc.add_paragraph()
    if indent_first:
        p.paragraph_format.first_line_indent = Cm(1.27)
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)
    return p


def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)
    return p


def add_number(text):
    p = doc.add_paragraph(style='List Number')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_runs(p, text)
    return p


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.27)
    p.paragraph_format.right_indent = Cm(1.27)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run('"' + text.strip('"').strip() + '"')
    run.italic = True
    return p


def add_table_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True


def add_figure_placeholder():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("[ tempat gambar ]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def add_figure_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = True


def add_figure_description(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    add_runs(p, text)


def parse_table_row(line):
    """Split markdown table row into cells, stripping outer pipes."""
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [c.strip() for c in line.split('|')]


def is_table_separator(line):
    """Check if line is a markdown table separator (e.g., | --- | --- |)."""
    s = line.strip()
    if not s.startswith('|'):
        return False
    cells = parse_table_row(s)
    return all(re.match(r'^:?-+:?$', c) for c in cells)


def add_table(rows):
    """Render parsed rows as a Word table.
    First row is treated as header (bold)."""
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Table Grid'
    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row[c_idx] if c_idx < len(row) else ''
            cell = table.cell(r_idx, c_idx)
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.line_spacing = 1.0
            add_runs(p, cell_text)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True


# ----- Parse markdown -----
with open(SRC, encoding='utf-8') as f:
    lines = f.read().splitlines()

i = 0
N = len(lines)
while i < N:
    raw = lines[i]
    line = raw.rstrip()

    # Horizontal rule
    if line.strip() == '---':
        i += 1
        continue

    # Empty line
    if not line.strip():
        i += 1
        continue

    # Heading
    m = re.match(r"^(#{1,6})\s+(.*)$", line)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        add_heading(text, level)
        i += 1
        continue

    # Table caption block: **Tabel X — ...**
    cap = re.match(r"^\*\*(Tabel\s+[^\*]+)\*\*$", line.strip())
    if cap and i + 1 < N and lines[i + 1].lstrip().startswith('|'):
        add_table_caption(cap.group(1).strip())
        i += 1
        continue

    # Figure block: **[Gambar X — ...]** optionally followed by *desc*
    fig = re.match(r"^\*\*\[(Gambar\s+[^\]]+)\]\*\*$", line.strip())
    if fig:
        add_figure_placeholder()
        add_figure_caption(fig.group(1).strip())
        if i + 1 < N:
            nxt = lines[i + 1].strip()
            m_desc = re.match(r"^\*([^*].*[^*])\*$", nxt)
            if m_desc and not nxt.startswith('**'):
                add_figure_description(m_desc.group(1))
                i += 2
                continue
        i += 1
        continue

    # Markdown table
    if line.lstrip().startswith('|'):
        rows = []
        j = i
        while j < N and lines[j].lstrip().startswith('|'):
            if is_table_separator(lines[j]):
                j += 1
                continue
            rows.append(parse_table_row(lines[j]))
            j += 1
        add_table(rows)
        i = j
        continue

    # Blockquote
    if line.lstrip().startswith('>'):
        buf = []
        j = i
        while j < N and lines[j].lstrip().startswith('>'):
            buf.append(lines[j].lstrip()[1:].strip())
            j += 1
        add_quote(' '.join(buf))
        i = j
        continue

    # Bullet list
    if re.match(r"^\s*-\s+", line):
        text = re.sub(r"^\s*-\s+", "", line)
        add_bullet(text)
        i += 1
        continue

    # Numbered list
    if re.match(r"^\s*\d+\.\s+", line):
        text = re.sub(r"^\s*\d+\.\s+", "", line)
        add_number(text)
        i += 1
        continue

    # Regular paragraph
    buf = [line]
    j = i + 1
    while j < N:
        nxt = lines[j].rstrip()
        if not nxt.strip():
            break
        if re.match(r"^(#{1,6})\s+", nxt):
            break
        if nxt.lstrip().startswith('>'):
            break
        if nxt.lstrip().startswith('|'):
            break
        if re.match(r"^\s*-\s+", nxt):
            break
        if re.match(r"^\s*\d+\.\s+", nxt):
            break
        if nxt.strip() == '---':
            break
        if re.match(r"^\*\*Tabel\s+", nxt.strip()):
            break
        if re.match(r"^\*\*\[Gambar\s+", nxt.strip()):
            break
        buf.append(nxt)
        j += 1
    add_paragraph(' '.join(buf))
    i = j

doc.save(DST)
print(f"Saved: {DST}")
