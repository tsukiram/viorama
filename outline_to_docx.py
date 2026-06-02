"""Convert outline_revisi8.md to properly formatted outline docx."""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = sys.argv[1] if len(sys.argv) >= 2 else "thesis/outline_revisi8.md"
DST = sys.argv[2] if len(sys.argv) >= 3 else "thesis/outline_revisi8.docx"

doc = Document()

# Page margins (4-3-3-3 Indonesian thesis)
section = doc.sections[0]
section.top_margin    = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin   = Cm(4)
section.right_margin  = Cm(3)

# Default style
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_after  = Pt(0)


def add_runs(p, text):
    """Parse inline *italic* and **bold** into runs."""
    pattern = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*([^*]+?)\*")
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        if m.group(1):
            r = p.add_run(m.group(1)); r.bold = True; r.italic = True
        elif m.group(2):
            r = p.add_run(m.group(2)); r.bold = True
        elif m.group(3):
            r = p.add_run(m.group(3)); r.italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)


def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(18)
    add_runs(p, text)


def add_bab(text):
    """BAB I, BAB II, etc."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(13)


# Indentation per level (in cm)
INDENT = {
    1: Cm(0),      # A. B. C.
    2: Cm(1.0),    # 1. 2. 3.
    3: Cm(2.0),    # a. b. c.
    4: Cm(3.0),    # 1) 2) 3)
    5: Cm(4.0),    # a) b) c)
}

# Patterns to detect outline level
LEVEL_PATTERNS = [
    (1, re.compile(r"^([A-Z])\.\s+(.+)$")),          # A. text
    (2, re.compile(r"^(\d+)\.\s+(.+)$")),             # 1. text
    (3, re.compile(r"^([a-z])\.\s+(.+)$")),           # a. text
    (4, re.compile(r"^(\d+)\)\s+(.+)$")),             # 1) text
    (5, re.compile(r"^([a-z])\)\s+(.+)$")),           # a) text
]


def detect_level(stripped):
    for lvl, pat in LEVEL_PATTERNS:
        m = pat.match(stripped)
        if m:
            return lvl, m.group(1), m.group(2)
    return None, None, None


def add_outline_item(level, marker, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = INDENT[level]
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(2)
    p.paragraph_format.line_spacing  = 1.15

    # Format marker
    if level == 1:
        full_marker = f"{marker}."
        r = p.add_run(f"{full_marker}\t")
        r.bold = True
        add_runs(p, text)
        for run in p.runs:
            if run != p.runs[0]:
                run.bold = True
    elif level == 2:
        full_marker = f"{marker}."
        p.add_run(f"{full_marker}\t")
        add_runs(p, text)
    elif level == 3:
        full_marker = f"{marker}."
        p.add_run(f"{full_marker}\t")
        add_runs(p, text)
    elif level == 4:
        full_marker = f"{marker})"
        p.add_run(f"{full_marker}\t")
        add_runs(p, text)
    elif level == 5:
        full_marker = f"{marker})"
        p.add_run(f"{full_marker}\t")
        add_runs(p, text)

    # Tab stop for alignment
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement('w:tabs')
    tab = OxmlElement('w:tab')
    tab.set(qn('w:val'), 'left')
    tab.set(qn('w:pos'), str(int(INDENT[level].pt * 20 + 360)))
    tabs.append(tab)
    pPr.append(tabs)


# ----- Parse -----
with open(SRC, encoding='utf-8') as f:
    raw_lines = f.read().splitlines()

for raw in raw_lines:
    line = raw.rstrip()

    # Skip blank and horizontal rules
    if not line.strip() or line.strip() == '---':
        continue

    # H1: title
    m1 = re.match(r"^#\s+(.+)$", line)
    if m1:
        add_title(m1.group(1))
        continue

    # H2: subtitle or BAB
    m2 = re.match(r"^##\s+(.+)$", line)
    if m2:
        text = m2.group(1).strip()
        if text.upper().startswith("BAB"):
            add_bab(text)
        else:
            add_subtitle(text)
        continue

    # Indented outline item — strip leading spaces
    stripped = line.strip()
    level, marker, text = detect_level(stripped)
    if level:
        add_outline_item(level, marker, text)
        continue

doc.save(DST)
print(f"Saved: {DST}")
